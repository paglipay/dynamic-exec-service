from __future__ import annotations

import io
from pathlib import Path

import pytest

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

from plugins.system_tools.file_reader_plugin import FileReaderPlugin


def _make_docx_file(tmp_path: Path, paragraphs: list[str], tables: list[list[list[str]]]) -> Path:
    """Build a DOCX file in tmp_path with given paragraphs and tables."""
    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)
    for table_data in tables:
        if not table_data:
            continue
        col_count = max(len(row) for row in table_data)
        tbl = doc.add_table(rows=len(table_data), cols=col_count)
        for r_idx, row_data in enumerate(table_data):
            for c_idx, cell_text in enumerate(row_data):
                tbl.rows[r_idx].cells[c_idx].text = cell_text
    file_path = tmp_path / "test.docx"
    doc.save(file_path)
    return file_path


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
def test_extract_docx_text_reads_paragraphs(tmp_path) -> None:
    docx_path = _make_docx_file(tmp_path, ["Hello World", "Second paragraph"], [])
    plugin = FileReaderPlugin(base_dir=str(tmp_path), allow_outside_base_dir=False)
    result = plugin.read_docx_text(str(docx_path))
    extracted = result["text"]
    assert "Hello World" in extracted
    assert "Second paragraph" in extracted


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
def test_extract_docx_text_reads_table_cells(tmp_path) -> None:
    table_data = [
        ["Site", "Loc Code", "School Name", "City", "Contractor"],
        ["MANN", "UCLA COMM SCH7574", "Horace Mann UCLA Community School", "LOS ANGELES, CA", "CSIB"],
    ]
    docx_path = _make_docx_file(tmp_path, [], [table_data])
    plugin = FileReaderPlugin(base_dir=str(tmp_path), allow_outside_base_dir=False)
    result = plugin.read_docx_text(str(docx_path))
    extracted = result["text"]

    assert "Site" in extracted
    assert "UCLA COMM SCH7574" in extracted
    assert "LOS ANGELES, CA" in extracted
    assert "CSIB" in extracted
    assert "Horace Mann UCLA Community School\tLOS ANGELES, CA" in extracted


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
def test_extract_docx_text_reads_paragraphs_and_tables_together(tmp_path) -> None:
    docx_path = _make_docx_file(
        tmp_path,
        ["CCTV Project - PUNCH LIST"],
        [[["Item", "Status"], ["Camera 1", "Done"], ["Camera 2", "Pending"]]],
    )
    plugin = FileReaderPlugin(base_dir=str(tmp_path), allow_outside_base_dir=False)
    result = plugin.read_docx_text(str(docx_path))
    extracted = result["text"]

    assert "CCTV Project - PUNCH LIST" in extracted
    assert "Camera 1" in extracted
    assert "Pending" in extracted


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
def test_extract_docx_text_raises_for_missing_file(tmp_path) -> None:
    plugin = FileReaderPlugin(base_dir=str(tmp_path), allow_outside_base_dir=False)
    with pytest.raises(FileNotFoundError):
        plugin.read_docx_text(str(tmp_path / "nonexistent.docx"))
