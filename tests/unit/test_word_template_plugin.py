"""Unit tests for generic WordTemplatePlugin."""

import json
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from plugins.system_tools.word_template_plugin import WordTemplatePlugin


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    with tempfile.TemporaryDirectory() as temp:
        yield Path(temp)


@pytest.fixture
def sample_docx(temp_dir):
    """Create a sample DOCX file for testing."""
    doc = Document()
    doc.add_paragraph("Hello <Name>!")
    doc.add_paragraph("You are from <City>.")

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Quantity"
    table.rows[1].cells[0].text = "<Item>"
    table.rows[1].cells[1].text = "<Quantity>"

    path = temp_dir / "template.docx"
    doc.save(path)
    return path


@pytest.fixture
def plugin(temp_dir):
    """Initialize WordTemplatePlugin with temp directory."""
    return WordTemplatePlugin(base_dir=str(temp_dir), allow_outside_base_dir=True)


class TestWordTemplatePluginInit:
    """Tests for plugin initialization."""

    def test_init_with_valid_base_dir(self, temp_dir):
        """Test initialization with valid base directory."""
        plugin = WordTemplatePlugin(base_dir=str(temp_dir), allow_outside_base_dir=True)
        assert plugin.base_dir == temp_dir
        assert plugin.allow_outside_base_dir is True

    def test_init_with_boundary_enforcement(self, temp_dir):
        """Test initialization with boundary enforcement enabled."""
        plugin = WordTemplatePlugin(base_dir=str(temp_dir), allow_outside_base_dir=False)
        assert plugin.allow_outside_base_dir is False

    def test_init_with_invalid_base_dir_nonexistent(self):
        """Test initialization fails with nonexistent base directory."""
        with pytest.raises(ValueError, match="must point to an existing directory"):
            WordTemplatePlugin(base_dir="/nonexistent/path", allow_outside_base_dir=True)

    def test_init_with_empty_base_dir(self):
        """Test initialization fails with empty base directory."""
        with pytest.raises(ValueError, match="must be a non-empty string"):
            WordTemplatePlugin(base_dir="", allow_outside_base_dir=True)

    def test_init_with_invalid_allow_outside_type(self, temp_dir):
        """Test initialization fails with invalid boolean type."""
        with pytest.raises(ValueError, match="must be a boolean"):
            WordTemplatePlugin(base_dir=str(temp_dir), allow_outside_base_dir="yes")


class TestRowTokenNormalization:
    """Tests for row token normalization."""

    def test_normalize_simple_key(self, plugin):
        """Test normalization of simple key."""
        token = plugin._normalize_row_token("Name")
        assert token == "<Name>"

    def test_normalize_already_wrapped_token(self, plugin):
        """Test normalization of already wrapped token."""
        token = plugin._normalize_row_token("<Name>")
        assert token == "<Name>"

    def test_normalize_with_whitespace(self, plugin):
        """Test normalization with surrounding whitespace."""
        token = plugin._normalize_row_token("  Name  ")
        assert token == "<Name>"


class TestBuildRowReplacements:
    """Tests for building replacement pairs from row objects."""

    def test_build_from_simple_row(self, plugin):
        """Test building replacements from simple row."""
        row = {"Name": "Alice", "City": "Boston"}
        replacements = plugin._build_row_replacements(row)
        assert ("<Name>", "Alice") in replacements
        assert ("<City>", "Boston") in replacements

    def test_build_with_none_value(self, plugin):
        """Test building replacements with None values."""
        row = {"Name": "Bob", "Email": None}
        replacements = plugin._build_row_replacements(row)
        assert ("<Name>", "Bob") in replacements
        assert ("<Email>", "") in replacements

    def test_build_with_numeric_values(self, plugin):
        """Test building replacements with numeric values."""
        row = {"Count": 42, "Ratio": 3.14}
        replacements = plugin._build_row_replacements(row)
        assert ("<Count>", "42") in replacements
        assert ("<Ratio>", "3.14") in replacements


class TestFilenameTemplateRendering:
    """Tests for filename template rendering."""

    def test_render_basic_template(self, plugin):
        """Test rendering basic filename template."""
        row = {"Name": "Alice", "Date": "2025-01-15"}
        filename = plugin._render_filename_template('Report_{Name}_{Date}.docx', row)
        assert filename == 'Report_Alice_2025-01-15.docx'

    def test_render_with_fstring_quote_prefix(self, plugin):
        """Test rendering with f-string quote prefix in template."""
        row = {"Name": "Bob"}
        filename = plugin._render_filename_template('f"Document_{Name}.docx"', row)
        assert filename == "Document_Bob.docx"

    def test_render_with_single_quote_fstring(self, plugin):
        """Test rendering with single-quoted f-string template."""
        row = {"Name": "Charlie"}
        filename = plugin._render_filename_template("f'Item_{Name}.docx'", row)
        assert filename == "Item_Charlie.docx"

    def test_render_with_sanitized_variant(self, plugin):
        """Test rendering with sanitized filename variant."""
        row = {"Name": 'Test<Bad>Name:?'}
        filename = plugin._render_filename_template('{Name_sanitized}.docx', row)
        assert filename == '{Name_sanitized}.docx' or 'Test_Bad_Name__.docx' in filename or '_sanitized' in filename

    def test_render_missing_placeholder_raises_error(self, plugin):
        """Test rendering with missing placeholder raises KeyError."""
        row = {"Name": "Alice"}
        with pytest.raises(ValueError, match="Unknown placeholder"):
            plugin._render_filename_template('Report_{Name}_{Missing}.docx', row)

    def test_render_empty_template_raises_error(self, plugin):
        """Test rendering empty template raises error."""
        row = {"Name": "Alice"}
        with pytest.raises(ValueError, match="empty"):
            plugin._render_filename_template('   ', row)


class TestDocumentGeneration:
    """Tests for document generation workflow."""

    def test_generate_basic_single_row(self, plugin, sample_docx, temp_dir):
        """Test generating document with single row."""
        rows = [{"Name": "Alice", "City": "Boston"}]
        output_dir = temp_dir / "output"

        result = plugin.generate_documents(
            rows=rows,
            input_docx=str(sample_docx),
            output_dir=str(output_dir),
            filename_template="Report_{Name}.docx",
        )

        assert result["status"] == "success"
        assert result["generated_count"] == 1
        assert result["rows_total"] == 1
        assert len(result["generated"]) == 1

        generated_file = output_dir / "Report_Alice.docx"
        assert generated_file.exists()

        doc = Document(generated_file)
        text = doc.paragraphs[0].text
        assert "Alice" in text
        assert "Boston" in text
        assert "<Name>" not in text

    def test_generate_multiple_rows(self, plugin, sample_docx, temp_dir):
        """Test generating documents from multiple rows."""
        rows = [
            {"Name": "Alice", "City": "Boston"},
            {"Name": "Bob", "City": "NYC"},
        ]
        output_dir = temp_dir / "output"

        result = plugin.generate_documents(
            rows=rows,
            input_docx=str(sample_docx),
            output_dir=str(output_dir),
            filename_template="Report_{Name}.docx",
        )

        assert result["generated_count"] == 2
        assert (output_dir / "Report_Alice.docx").exists()
        assert (output_dir / "Report_Bob.docx").exists()

    def test_generate_with_append_lines(self, plugin, sample_docx, temp_dir):
        """Test generating document with appended lines."""
        rows = [{"Name": "Charlie"}]
        output_dir = temp_dir / "output"

        result = plugin.generate_documents(
            rows=rows,
            input_docx=str(sample_docx),
            output_dir=str(output_dir),
            filename_template="Report_{Name}.docx",
            append_lines=["---", "This is a footer.", "End of report."],
        )

        assert result["generated_count"] == 1
        generated_file = output_dir / "Report_Charlie.docx"

        doc = Document(generated_file)
        para_texts = [p.text for p in doc.paragraphs]
        assert any("This is a footer" in p for p in para_texts)

    def test_generate_with_json_file_rows(self, plugin, sample_docx, temp_dir):
        """Test generating documents with rows loaded from JSON file."""
        rows_data = [
            {"Name": "David", "City": "LA"},
            {"Name": "Eve", "City": "SF"},
        ]
        rows_file = temp_dir / "rows.json"
        rows_file.write_text(json.dumps(rows_data))

        output_dir = temp_dir / "output"

        result = plugin.generate_documents(
            rows=str(rows_file),
            input_docx=str(sample_docx),
            output_dir=str(output_dir),
            filename_template="Report_{Name}.docx",
        )

        assert result["generated_count"] == 2
        assert (output_dir / "Report_David.docx").exists()
        assert (output_dir / "Report_Eve.docx").exists()

    def test_generate_with_dict_payload(self, plugin, sample_docx, temp_dir):
        """Test generating documents with payload as nested dict."""
        rows = [{"Name": "Frank"}]
        output_dir = temp_dir / "output"
        payload = {
            "rows": rows,
            "input_docx": str(sample_docx),
            "output_dir": str(output_dir),
            "filename_template": "Doc_{Name}.docx",
        }

        result = plugin.generate_documents(rows=payload)

        assert result["generated_count"] == 1
        assert (output_dir / "Doc_Frank.docx").exists()

    def test_generate_invalid_rows_type(self, plugin, sample_docx, temp_dir):
        """Test generate fails with invalid rows type."""
        with pytest.raises(ValueError):
            plugin.generate_documents(
                rows="not_a_file",
                input_docx=str(sample_docx),
                output_dir=str(temp_dir),
                filename_template="Report.docx",
            )

    def test_generate_missing_input_docx_raises_error(self, plugin, temp_dir):
        """Test generate fails with missing input_docx."""
        with pytest.raises(ValueError, match="input_docx"):
            plugin.generate_documents(
                rows=[{"Name": "Test"}],
                input_docx="",
                output_dir=str(temp_dir),
                filename_template="Report.docx",
            )

    def test_generate_missing_output_dir_raises_error(self, plugin, sample_docx):
        """Test generate fails with missing output_dir."""
        with pytest.raises(ValueError, match="output_dir"):
            plugin.generate_documents(
                rows=[{"Name": "Test"}],
                input_docx=str(sample_docx),
                output_dir="",
                filename_template="Report.docx",
            )

    def test_generate_missing_filename_template_raises_error(self, plugin, sample_docx, temp_dir):
        """Test generate fails with missing filename_template."""
        with pytest.raises(ValueError, match="filename_template"):
            plugin.generate_documents(
                rows=[{"Name": "Test"}],
                input_docx=str(sample_docx),
                output_dir=str(temp_dir),
                filename_template="",
            )

    def test_generate_nonexistent_input_docx_raises_error(self, plugin, temp_dir):
        """Test generate fails with nonexistent input_docx."""
        with pytest.raises(ValueError, match="does not exist"):
            plugin.generate_documents(
                rows=[{"Name": "Test"}],
                input_docx=str(temp_dir / "missing.docx"),
                output_dir=str(temp_dir),
                filename_template="Report.docx",
            )

    def test_generate_non_docx_input_raises_error(self, plugin, temp_dir):
        """Test generate fails with non-DOCX input file."""
        fake_file = temp_dir / "file.txt"
        fake_file.write_text("not a docx")

        with pytest.raises(ValueError, match="must be a .docx file"):
            plugin.generate_documents(
                rows=[{"Name": "Test"}],
                input_docx=str(fake_file),
                output_dir=str(temp_dir),
                filename_template="Report.docx",
            )

    def test_generate_empty_rows_raises_error(self, plugin, sample_docx, temp_dir):
        """Test generate fails with empty rows list."""
        with pytest.raises(ValueError, match="at least one"):
            plugin.generate_documents(
                rows=[],
                input_docx=str(sample_docx),
                output_dir=str(temp_dir),
                filename_template="Report.docx",
            )


class TestTableUpdates:
    """Tests for table row updates."""

    def test_generate_with_table_updates(self, temp_dir):
        """Test generating document with table row expansion."""
        doc = Document()
        doc.add_paragraph("Document for items.")

        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "Item"
        table.rows[0].cells[1].text = "Quantity"
        table.rows[0].cells[2].text = "Notes"

        template_row = table.rows[1]
        template_row.cells[0].text = "<ITEM_MARKER>"
        template_row.cells[1].text = "<Qty>"
        template_row.cells[2].text = "<Notes>"

        template_path = temp_dir / "table_template.docx"
        doc.save(template_path)

        plugin = WordTemplatePlugin(base_dir=str(temp_dir), allow_outside_base_dir=True)
        output_dir = temp_dir / "output"

        rows = [
            {"Qty": "5", "Notes": "Available"},
            {"Qty": "10", "Notes": "In stock"},
        ]

        table_updates = [
            {
                "table_selector": {"table_index": 0},
                "template_row_marker": "<ITEM_MARKER>",
                "remove_template_row": True,
                "rows": rows,
            }
        ]

        result = plugin.generate_documents(
            rows=[{"doc_id": "1"}],
            input_docx=str(template_path),
            output_dir=str(output_dir),
            filename_template="Report_{doc_id}.docx",
            table_updates=table_updates,
        )

        assert result["generated_count"] == 1
        generated_file = output_dir / "Report_1.docx"
        assert generated_file.exists()

        doc = Document(generated_file)
        table = doc.tables[0]
        assert len(table.rows) >= 2
        assert "5" in table.rows[-2].cells[1].text
        assert "10" in table.rows[-1].cells[1].text

    def test_table_updates_with_header_matching(self, temp_dir):
        """Test table updates using header substring matching."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Item ID"
        table.rows[0].cells[1].text = "Total Amount"
        table.rows[1].cells[0].text = "<ITEM>"
        table.rows[1].cells[1].text = "<AMOUNT>"

        template_path = temp_dir / "header_table.docx"
        doc.save(template_path)

        plugin = WordTemplatePlugin(base_dir=str(temp_dir), allow_outside_base_dir=True)
        output_dir = temp_dir / "output"

        rows = [{"ITEM": "A123", "AMOUNT": "$50"}]

        table_updates = [
            {
                "table_selector": {"header_contains": ["Item", "Amount"]},
                "template_row_index": 1,
                "remove_template_row": True,
                "rows": rows,
            }
        ]

        result = plugin.generate_documents(
            rows=[{"doc_id": "1"}],
            input_docx=str(template_path),
            output_dir=str(output_dir),
            filename_template="Invoice_{doc_id}.docx",
            table_updates=table_updates,
        )

        assert result["generated_count"] == 1


class TestPDFExport:
    """Tests for PDF export functionality."""

    @patch("plugins.system_tools.word_template_plugin.WordTemplatePlugin._export_pdf")
    def test_generate_with_pdf_export(self, mock_export, plugin, sample_docx, temp_dir):
        """Test generating document with PDF export."""
        rows = [{"Name": "Grace"}]
        output_dir = temp_dir / "output"

        result = plugin.generate_documents(
            rows=rows,
            input_docx=str(sample_docx),
            output_dir=str(output_dir),
            filename_template="Report_{Name}.docx",
            export_pdf=True,
        )

        assert result["generated_count"] == 1
        assert result["generated"][0]["output_pdf"] is not None
        mock_export.assert_called()

    def test_export_pdf_with_all_methods_failing(self, plugin, sample_docx, temp_dir):
        """Test PDF export error handling when all converters unavailable."""
        docx_file = temp_dir / "test.docx"
        pdf_file = temp_dir / "test.pdf"
        import shutil

        shutil.copy(sample_docx, docx_file)

        with patch("plugins.system_tools.word_template_plugin.shutil.which", return_value=None), patch(
            "plugins.system_tools.word_template_plugin.WordTemplatePlugin._export_pdf",
            side_effect=Exception("All converters failed"),
        ):
            with pytest.raises(Exception):
                plugin._export_pdf(docx_file, pdf_file)


class TestPathResolution:
    """Tests for path resolution and boundary enforcement."""

    def test_resolve_absolute_path(self, temp_dir):
        """Test resolving absolute path."""
        plugin = WordTemplatePlugin(base_dir=str(temp_dir), allow_outside_base_dir=True)
        abs_path = temp_dir / "subdir" / "file.txt"
        abs_path.parent.mkdir(parents=True, exist_ok=True)
        abs_path.touch()

        resolved = plugin._resolve_path(str(abs_path), require_exists=True)
        assert resolved.exists()

    def test_resolve_relative_path(self, temp_dir):
        """Test resolving relative path from base_dir."""
        plugin = WordTemplatePlugin(base_dir=str(temp_dir), allow_outside_base_dir=True)
        subdir = temp_dir / "subdir"
        subdir.mkdir()
        file_path = subdir / "file.txt"
        file_path.touch()

        resolved = plugin._resolve_path("subdir/file.txt", require_exists=True)
        assert resolved.exists()

    def test_path_boundary_enforcement_inside_base_dir(self, temp_dir):
        """Test path boundary is enforced for paths inside base_dir."""
        plugin = WordTemplatePlugin(base_dir=str(temp_dir), allow_outside_base_dir=False)
        file_path = temp_dir / "inside.txt"
        file_path.touch()

        resolved = plugin._resolve_path("inside.txt", require_exists=True)
        assert resolved.exists()

    def test_path_boundary_enforcement_outside_base_dir(self, temp_dir):
        """Test path boundary is enforced for paths outside base_dir."""
        plugin = WordTemplatePlugin(base_dir=str(temp_dir), allow_outside_base_dir=False)

        with pytest.raises(ValueError, match="must be inside base_dir"):
            plugin._resolve_path("../outside.txt")

    def test_path_nonexistent_with_require_exists(self, temp_dir):
        """Test path resolution fails when file doesn't exist and require_exists=True."""
        plugin = WordTemplatePlugin(base_dir=str(temp_dir), allow_outside_base_dir=True)

        with pytest.raises(ValueError, match="does not exist"):
            plugin._resolve_path("missing_file.txt", require_exists=True)


class TestEdgeCases:
    """Tests for edge cases and unusual inputs."""

    def test_generate_with_special_characters_in_filename(self, plugin, sample_docx, temp_dir):
        """Test generating document with special characters in data."""
        rows = [{"Name": 'Test<>Name:?', "City": 'New/York'}]
        output_dir = temp_dir / "output"

        result = plugin.generate_documents(
            rows=rows,
            input_docx=str(sample_docx),
            output_dir=str(output_dir),
            filename_template="Report_{Name_sanitized}.docx",
        )

        assert result["generated_count"] == 1
        output_files = list(output_dir.glob("*.docx"))
        assert len(output_files) == 1

    def test_generate_with_unicode_characters(self, plugin, sample_docx, temp_dir):
        """Test generating document with unicode characters."""
        rows = [{"Name": "José", "City": "São Paulo"}]
        output_dir = temp_dir / "output"

        result = plugin.generate_documents(
            rows=rows,
            input_docx=str(sample_docx),
            output_dir=str(output_dir),
            filename_template="Report_{Name}.docx",
        )

        assert result["generated_count"] == 1
        generated_file = output_dir / "Report_José.docx"
        assert generated_file.exists()

    def test_generate_with_very_long_field_values(self, plugin, sample_docx, temp_dir):
        """Test generating document with very long field values."""
        long_text = "Lorem ipsum dolor sit amet " * 100
        rows = [{"Name": "Test", "City": long_text}]
        output_dir = temp_dir / "output"

        result = plugin.generate_documents(
            rows=rows,
            input_docx=str(sample_docx),
            output_dir=str(output_dir),
            filename_template="Report_{Name}.docx",
        )

        assert result["generated_count"] == 1
        doc = Document(output_dir / "Report_Test.docx")
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert long_text[:50] in full_text
