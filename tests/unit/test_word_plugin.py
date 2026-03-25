from __future__ import annotations

import json

from docx import Document
import pytest

from plugins.system_tools.word_plugin import WordPlugin


def test_create_and_inspect_document(tmp_path) -> None:
    plugin = WordPlugin(base_dir=str(tmp_path), allow_outside_base_dir=False)

    create_result = plugin.create_document(
        {
            "title": "Quarterly Update",
            "paragraphs": ["Line one", "Line two"],
            "output_path": "report.docx",
        }
    )

    assert create_result["status"] == "success"
    assert create_result["paragraph_count"] == 2

    inspect_result = plugin.inspect_document(
        {
            "file_path": "report.docx",
            "max_paragraphs": 5,
        }
    )

    assert inspect_result["status"] == "success"
    assert inspect_result["paragraph_count"] == 3
    assert inspect_result["paragraph_preview"] == ["Quarterly Update", "Line one", "Line two"]
    assert inspect_result["table_count"] == 0


def test_replace_text_updates_docx_output(tmp_path) -> None:
    source_path = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("Hello <NAME>")
    document.save(source_path)

    plugin = WordPlugin(base_dir=str(tmp_path), allow_outside_base_dir=False)
    result = plugin.replace_text(
        {
            "file_path": "template.docx",
            "replacements": [{"find": "<NAME>", "replace": "Ada"}],
            "output_path": "filled.docx",
        }
    )

    assert result["status"] == "success"
    assert result["changed_block_count"] == 1

    updated = Document(tmp_path / "filled.docx")
    assert updated.paragraphs[0].text == "Hello Ada"


def test_add_table_appends_rows(tmp_path) -> None:
    source_path = tmp_path / "base.docx"
    document = Document()
    document.add_paragraph("Inventory")
    document.save(source_path)

    plugin = WordPlugin(base_dir=str(tmp_path), allow_outside_base_dir=False)
    result = plugin.add_table(
        {
            "file_path": "base.docx",
            "headers": ["Item", "Qty"],
            "rows": [["Pens", 3], ["Paper", 10]],
            "output_path": "with_table.docx",
        }
    )

    assert result["status"] == "success"
    assert result["row_count"] == 2

    updated = Document(tmp_path / "with_table.docx")
    assert len(updated.tables) == 1
    assert updated.tables[0].rows[0].cells[0].text == "Item"
    assert updated.tables[0].rows[1].cells[0].text == "Pens"
    assert updated.tables[0].rows[2].cells[1].text == "10"


def test_generate_documents_fills_template_and_expands_table_rows(tmp_path) -> None:
    template_path = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("Client: <client_name>")
    document.add_paragraph("Project: <project_code>")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "<ROW_TEMPLATE><item>"
    table.rows[1].cells[1].text = "<value>"
    document.save(template_path)

    plugin = WordPlugin(base_dir=str(tmp_path), allow_outside_base_dir=False)
    result = plugin.generate_documents(
        {
            "documents_json": [
                {
                    "document_name": "North Summary",
                    "client_name": "Contoso",
                    "project_code": "NR-100",
                    "owner": "Alice",
                }
            ],
            "input_docx": "template.docx",
            "output_dir": "out",
            "filename_template": "{document_name_sanitized}.docx",
            "append_lines": ["Generated automatically"],
            "table_updates": [
                {
                    "table_selector": {"table_index": 0},
                    "template_row_marker": "<ROW_TEMPLATE>",
                    "rows": [
                        {"item": "Owner", "value": "<owner>"},
                        {"item": "Code", "value": "<project_code>"},
                    ],
                }
            ],
        }
    )

    assert result["status"] == "success"
    assert result["generated_count"] == 1

    output_docx = tmp_path / "out" / "North Summary.docx"
    assert output_docx.exists()

    generated = Document(output_docx)
    paragraph_texts = [paragraph.text for paragraph in generated.paragraphs]
    assert "Client: Contoso" in paragraph_texts
    assert "Project: NR-100" in paragraph_texts
    assert "Generated automatically" in paragraph_texts
    assert len(generated.tables) == 1
    assert len(generated.tables[0].rows) == 3
    assert generated.tables[0].rows[1].cells[0].text == "Owner"
    assert generated.tables[0].rows[1].cells[1].text == "Alice"
    assert generated.tables[0].rows[2].cells[0].text == "Code"
    assert generated.tables[0].rows[2].cells[1].text == "NR-100"


def test_generate_documents_accepts_json_files(tmp_path) -> None:
    template_path = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("Hello <document_name>")
    document.save(template_path)

    documents_json_path = tmp_path / "documents.json"
    documents_json_path.write_text(
        json.dumps([{"document_name": "Brief"}]),
        encoding="utf-8",
    )

    plugin = WordPlugin(base_dir=str(tmp_path), allow_outside_base_dir=False)
    result = plugin.generate_documents(
        {
            "documents_json": "documents.json",
            "input_docx": "template.docx",
            "output_dir": "generated",
            "filename_template": "{document_name_sanitized}.docx",
        }
    )

    assert result["status"] == "success"
    assert (tmp_path / "generated" / "Brief.docx").exists()


def test_export_pdf_requires_available_converter(tmp_path) -> None:
    source_path = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("Export me")
    document.save(source_path)

    plugin = WordPlugin(base_dir=str(tmp_path), allow_outside_base_dir=False)

    with pytest.raises(ValueError, match="PDF export failed"):
        plugin.export_pdf({"file_path": "source.docx", "output_path": "source.pdf"})