"""Word template plugin for generic document generation from header-driven data rows."""

from __future__ import annotations

from copy import deepcopy
import json
import shutil
import subprocess
from pathlib import Path
from typing import Any


class WordTemplatePlugin:
    """Generate DOCX files from row data with exact header-to-token matching and optional table expansion."""

    def __init__(self, base_dir: str = "generated_data", allow_outside_base_dir: bool = True) -> None:
        """Initialize plugin with base directory and optional path boundary enforcement."""
        if not isinstance(base_dir, str) or not base_dir.strip():
            raise ValueError("base_dir must be a non-empty string")
        if not isinstance(allow_outside_base_dir, bool):
            raise ValueError("allow_outside_base_dir must be a boolean")

        self.base_dir = Path(base_dir).resolve()
        self.allow_outside_base_dir = allow_outside_base_dir

        if not self.base_dir.exists() or not self.base_dir.is_dir():
            raise ValueError("base_dir must point to an existing directory")

    def _resolve_path(self, path_value: str, require_exists: bool = False) -> Path:
        """Resolve a path relative to base_dir with optional boundary enforcement."""
        if not isinstance(path_value, str) or not path_value.strip():
            raise ValueError("path must be a non-empty string")

        raw = Path(path_value.strip())
        resolved = raw.resolve() if raw.is_absolute() else (self.base_dir / raw).resolve()

        if not self.allow_outside_base_dir:
            try:
                resolved.relative_to(self.base_dir)
            except ValueError as exc:
                raise ValueError("path must be inside base_dir") from exc

        if require_exists and not resolved.exists():
            raise ValueError(f"path does not exist: {resolved}")

        return resolved

    def _load_json_array(self, path: Path) -> list[Any]:
        """Load and validate a JSON file as a top-level array."""
        if not path.exists() or not path.is_file():
            raise ValueError(f"JSON file not found: {path}")

        try:
            with path.open("r", encoding="utf-8") as handle:
                data = json.load(handle)
        except Exception as exc:
            raise ValueError(f"Failed to read JSON file: {path}") from exc

        if not isinstance(data, list):
            raise ValueError(f"Expected top-level JSON array in file: {path}")

        return data

    def _load_document(self, path: Path) -> Any:
        """Load a DOCX file using python-docx."""
        try:
            from docx import Document  # type: ignore[import-not-found]
        except Exception as exc:
            raise ValueError("python-docx is not installed. Install requirements.txt dependencies.") from exc

        return Document(path)

    def _replace_in_paragraph_preserve_format(self, paragraph: Any, replacements: list[tuple[str, str]]) -> bool:
        """Replace tokens in a paragraph while preserving run-level formatting."""
        changed = False

        for run in paragraph.runs:
            original = run.text
            updated = original
            for old, new in replacements:
                updated = updated.replace(old, new)
            if updated != original:
                run.text = updated
                changed = True

        for old, new in replacements:
            if not old:
                continue

            while True:
                full_text = "".join(run.text for run in paragraph.runs)
                start = full_text.find(old)
                if start == -1:
                    break

                end = start + len(old)
                pos = 0
                start_run = -1
                start_offset = -1
                end_run = -1
                end_offset = -1

                for idx, run in enumerate(paragraph.runs):
                    run_len = len(run.text)
                    next_pos = pos + run_len

                    if start_run == -1 and start < next_pos:
                        start_run = idx
                        start_offset = start - pos

                    if end_run == -1 and end <= next_pos:
                        end_run = idx
                        end_offset = end - pos
                        break

                    pos = next_pos

                if start_run == -1 or end_run == -1:
                    break

                if start_run == end_run:
                    run = paragraph.runs[start_run]
                    text = run.text
                    run.text = text[:start_offset] + new + text[end_offset:]
                else:
                    first_run = paragraph.runs[start_run]
                    last_run = paragraph.runs[end_run]

                    prefix = first_run.text[:start_offset]
                    suffix = last_run.text[end_offset:]

                    first_run.text = prefix + new
                    for idx in range(start_run + 1, end_run):
                        paragraph.runs[idx].text = ""
                    last_run.text = suffix

                changed = True

        return changed

    def _replace_in_document(self, document: Any, replacements: list[tuple[str, str]]) -> int:
        """Replace tokens across all paragraphs and table cells in a document."""
        changed_count = 0

        for paragraph in document.paragraphs:
            if self._replace_in_paragraph_preserve_format(paragraph, replacements):
                changed_count += 1

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if self._replace_in_paragraph_preserve_format(paragraph, replacements):
                            changed_count += 1

        return changed_count

    def _append_paragraphs(self, document: Any, lines: list[str]) -> None:
        """Append lines as paragraphs to a document."""
        for line in lines:
            document.add_paragraph(line)

    def _sanitize_filename(self, name: str) -> str:
        """Sanitize a string for use as a filesystem filename."""
        invalid = '<>:"/\\|?*'
        sanitized = "".join("_" if char in invalid else char for char in name).strip()
        return sanitized or "document"

    def _export_pdf(self, docx_path: Path, pdf_path: Path) -> None:
        """Export DOCX to PDF using available local converters."""
        errors: list[str] = []

        try:
            from docx2pdf import convert  # type: ignore[import-not-found]

            convert(str(docx_path), str(pdf_path))
            return
        except Exception as exc:
            errors.append(f"docx2pdf failed: {exc}")

        try:
            import win32com.client  # type: ignore[import-not-found]

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            source = word.Documents.Open(str(docx_path.resolve()))
            source.SaveAs(str(pdf_path.resolve()), FileFormat=17)
            source.Close()
            word.Quit()
            return
        except Exception as exc:
            errors.append(f"win32com failed: {exc}")

        try:
            soffice_cmd = shutil.which("soffice")
            if not soffice_cmd:
                raise FileNotFoundError("'soffice' not found on PATH")

            output_dir = pdf_path.resolve().parent
            output_dir.mkdir(parents=True, exist_ok=True)

            subprocess.run(
                [
                    soffice_cmd,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_dir),
                    str(docx_path.resolve()),
                ],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )

            generated_pdf = output_dir / f"{docx_path.stem}.pdf"
            if not generated_pdf.exists():
                raise RuntimeError("LibreOffice conversion completed but PDF file was not created.")

            if generated_pdf.resolve() != pdf_path.resolve():
                shutil.move(str(generated_pdf), str(pdf_path.resolve()))
            return
        except Exception as exc:
            errors.append(f"LibreOffice/soffice failed: {exc}")

        details = "\n".join(f"- {item}" for item in errors)
        raise ValueError(
            "PDF export failed. Use one of these options:\n"
            "1) Install Microsoft Word + docx2pdf/pywin32, or\n"
            "2) Install LibreOffice and ensure 'soffice' is on PATH.\n\n"
            "Attempt details:\n"
            f"{details}"
        )

    def _ensure_array(self, value: Any, field_name: str) -> list[Any]:
        """Load data as array from direct list or JSON file path."""
        if isinstance(value, list):
            return value
        if isinstance(value, str) and value.strip():
            path = self._resolve_path(value, require_exists=True)
            if path.suffix.lower() != ".json":
                raise ValueError(f"{field_name} path must be a .json file")
            return self._load_json_array(path)
        raise ValueError(f"{field_name} must be a JSON file path or an array")

    def _normalize_row_token(self, key: str) -> str:
        """Convert a row key to a placeholder token (<Key>)."""
        token = str(key).strip()
        if token.startswith("<") and token.endswith(">") and len(token) >= 3:
            return token
        return f"<{token}>"

    def _build_row_replacements(self, row_item: dict[str, Any]) -> list[tuple[str, str]]:
        """Build (token, value) tuples from a row object using exact header names."""
        pairs: list[tuple[str, str]] = []
        for key, value in row_item.items():
            token = self._normalize_row_token(str(key))
            rendered = "" if value is None else str(value)
            pairs.append((token, rendered))
        return pairs

    def _render_filename_template(self, filename_template: str, row_context: dict[str, Any]) -> str:
        """Render output filename using row fields as format context."""
        stripped_template = filename_template.strip()
        if stripped_template.startswith('f"') and stripped_template.endswith('"'):
            stripped_template = stripped_template[2:-1]
        elif stripped_template.startswith("f'") and stripped_template.endswith("'"):
            stripped_template = stripped_template[2:-1]

        safe_context = {}
        for key, value in row_context.items():
            safe_context[key] = "" if value is None else str(value)
            safe_context[f"{key}_sanitized"] = self._sanitize_filename(str(value) if value is not None else "")

        try:
            rendered = stripped_template.format(**safe_context)
        except KeyError as exc:
            raise ValueError(f"Unknown placeholder in filename template: {exc}") from exc

        rendered = rendered.strip()
        if not rendered:
            raise ValueError("Rendered filename is empty.")
        return rendered

    def _find_table_index(self, document: Any, selector: dict[str, Any], update_index: int) -> int:
        """Find a table by explicit index or header substring match."""
        table_count = len(document.tables)
        if table_count <= 0:
            raise ValueError("No tables found in document")

        table_index_value = selector.get("table_index")
        if table_index_value is not None:
            if not isinstance(table_index_value, int):
                raise ValueError(f"table_updates[{update_index}].table_selector.table_index must be an integer")
            if table_index_value < 0 or table_index_value >= table_count:
                raise ValueError(
                    f"table_updates[{update_index}] table_index out of range: {table_index_value}; total_tables={table_count}"
                )
            return table_index_value

        header_contains = selector.get("header_contains")
        if not isinstance(header_contains, list) or not header_contains:
            raise ValueError(
                f"table_updates[{update_index}].table_selector must include table_index or non-empty header_contains"
            )

        normalized_terms = [str(term).strip().lower() for term in header_contains if str(term).strip()]
        if not normalized_terms:
            raise ValueError(
                f"table_updates[{update_index}].table_selector.header_contains must have non-empty string items"
            )

        for idx, table in enumerate(document.tables):
            if not table.rows:
                continue
            header_text = " | ".join(cell.text for cell in table.rows[0].cells).lower()
            if all(term in header_text for term in normalized_terms):
                return idx

        raise ValueError(
            f"table_updates[{update_index}] did not match any table using header_contains={normalized_terms}"
        )

    def _find_template_row_index(self, table: Any, marker: str | None, row_index: int | None, update_index: int) -> int:
        """Find a template row by explicit index or marker substring match."""
        if isinstance(row_index, int):
            if row_index < 0 or row_index >= len(table.rows):
                raise ValueError(
                    f"table_updates[{update_index}].template_row_index out of range: {row_index}; total_rows={len(table.rows)}"
                )
            return row_index

        if isinstance(marker, str) and marker.strip():
            marker_value = marker.strip()
            for idx, row in enumerate(table.rows):
                row_text = "\n".join(cell.text for cell in row.cells)
                if marker_value in row_text:
                    return idx
            raise ValueError(
                f"table_updates[{update_index}] could not find template row marker: {marker_value}"
            )

        raise ValueError(f"table_updates[{update_index}] requires template_row_marker or template_row_index")

    def _clear_row_marker(self, row: Any, marker: str) -> None:
        """Clear a marker token from a table row."""
        marker_value = marker.strip()
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if marker_value in paragraph.text:
                    self._replace_in_paragraph_preserve_format(paragraph, [(marker_value, "")])

    def _replace_tokens_in_row(self, row: Any, replacements: list[tuple[str, str]]) -> int:
        """Replace tokens in a table row's cells."""
        changed_count = 0
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if self._replace_in_paragraph_preserve_format(paragraph, replacements):
                    changed_count += 1
        return changed_count

    def _remove_row(self, table: Any, row_index: int) -> None:
        """Remove a row from a table."""
        row = table.rows[row_index]
        row_element = row._tr
        parent = row_element.getparent()
        parent.remove(row_element)

    def _apply_table_updates(
        self,
        document: Any,
        table_updates: list[Any],
        document_replacements: list[tuple[str, str]],
    ) -> tuple[int, int, int]:
        """Apply table updates: find table, clone template row for each data row, optionally remove template."""
        tables_updated = 0
        rows_created = 0
        changed_blocks = 0

        for update_index, update in enumerate(table_updates):
            if not isinstance(update, dict):
                raise ValueError(f"table_updates[{update_index}] must be an object")

            selector = update.get("table_selector")
            if not isinstance(selector, dict):
                raise ValueError(f"table_updates[{update_index}].table_selector must be an object")

            rows_data = update.get("rows")
            if not isinstance(rows_data, list):
                raise ValueError(f"table_updates[{update_index}].rows must be an array")

            table_index = self._find_table_index(document, selector, update_index)
            table = document.tables[table_index]

            template_row_marker = update.get("template_row_marker")
            template_row_index = update.get("template_row_index")
            resolved_template_row_index = self._find_template_row_index(
                table,
                template_row_marker if isinstance(template_row_marker, str) else None,
                template_row_index if isinstance(template_row_index, int) else None,
                update_index,
            )

            remove_template_row = update.get("remove_template_row", True)
            if not isinstance(remove_template_row, bool):
                raise ValueError(f"table_updates[{update_index}].remove_template_row must be a boolean")

            template_row = table.rows[resolved_template_row_index]
            if isinstance(template_row_marker, str) and template_row_marker.strip():
                self._clear_row_marker(template_row, template_row_marker)

            insert_after = template_row._tr
            created_for_update = 0
            for row_item_index, row_item in enumerate(rows_data):
                if not isinstance(row_item, dict):
                    raise ValueError(
                        f"table_updates[{update_index}].rows[{row_item_index}] must be an object of token/value pairs"
                    )

                new_tr = deepcopy(template_row._tr)
                insert_after.addnext(new_tr)
                insert_after = new_tr

                created_row = table.rows[resolved_template_row_index + created_for_update + 1]
                row_replacements = self._build_row_replacements(row_item)
                changed_blocks += self._replace_tokens_in_row(created_row, row_replacements)
                created_for_update += 1
                rows_created += 1

            if remove_template_row:
                self._remove_row(table, resolved_template_row_index)

            tables_updated += 1

        return tables_updated, rows_created, changed_blocks

    def generate_documents(
        self,
        rows: str | list[Any] | dict[str, Any],
        input_docx: str | None = None,
        output_dir: str | None = None,
        filename_template: str | None = None,
        append_lines: list[str] | None = None,
        table_updates: list[Any] | None = None,
        export_pdf: bool = False,
    ) -> dict[str, Any]:
        """Generate DOCX files from rows using exact header-to-token matching."""
        if isinstance(rows, dict):
            payload = rows
            resolved_rows = payload.get("rows")
            resolved_input_docx = payload.get("input_docx", input_docx)
            resolved_output_dir = payload.get("output_dir", output_dir)
            resolved_filename_template = payload.get("filename_template", filename_template)
            resolved_append_lines = payload.get("append_lines", append_lines)
            resolved_table_updates = payload.get("table_updates", table_updates)
            resolved_export_pdf = payload.get("export_pdf", export_pdf)
        else:
            resolved_rows = rows
            resolved_input_docx = input_docx
            resolved_output_dir = output_dir
            resolved_filename_template = filename_template
            resolved_append_lines = append_lines
            resolved_table_updates = table_updates
            resolved_export_pdf = export_pdf

        if not isinstance(resolved_input_docx, str) or not resolved_input_docx.strip():
            raise ValueError("input_docx is required")
        if not isinstance(resolved_output_dir, str) or not resolved_output_dir.strip():
            raise ValueError("output_dir is required")
        if not isinstance(resolved_filename_template, str) or not resolved_filename_template.strip():
            raise ValueError("filename_template is required")
        if not isinstance(resolved_export_pdf, bool):
            raise ValueError("export_pdf must be a boolean")
        if resolved_append_lines is not None and not isinstance(resolved_append_lines, list):
            raise ValueError("append_lines must be an array when provided")
        if resolved_table_updates is not None and not isinstance(resolved_table_updates, list):
            raise ValueError("table_updates must be an array when provided")

        rows_data = self._ensure_array(resolved_rows, "rows")
        normalized_rows = [item for item in rows_data if isinstance(item, dict)]

        if not normalized_rows:
            raise ValueError("rows must contain at least one object")

        safe_append_lines: list[str] = []
        if isinstance(resolved_append_lines, list):
            for line in resolved_append_lines:
                if isinstance(line, str) and line.strip():
                    safe_append_lines.append(line)

        input_path = self._resolve_path(resolved_input_docx, require_exists=True)
        output_path = self._resolve_path(resolved_output_dir, require_exists=False)

        if input_path.suffix.lower() != ".docx":
            raise ValueError("input_docx must be a .docx file")

        output_path.mkdir(parents=True, exist_ok=True)

        generated: list[dict[str, Any]] = []

        for row_item in normalized_rows:
            replacements = self._build_row_replacements(row_item)
            filename = self._render_filename_template(resolved_filename_template, row_item)
            output_docx_path = output_path / filename
            output_docx_path.parent.mkdir(parents=True, exist_ok=True)

            shutil.copy2(input_path, output_docx_path)

            document = self._load_document(output_docx_path)

            table_count_updated = 0
            rows_generated = 0
            table_changed_blocks = 0
            if isinstance(resolved_table_updates, list) and resolved_table_updates:
                table_count_updated, rows_generated, table_changed_blocks = self._apply_table_updates(
                    document=document,
                    table_updates=resolved_table_updates,
                    document_replacements=replacements,
                )

            changed_count = self._replace_in_document(document, replacements)
            changed_count += table_changed_blocks
            self._append_paragraphs(document, safe_append_lines)
            document.save(output_docx_path)

            output_pdf_path: str | None = None
            if resolved_export_pdf:
                pdf_path = output_docx_path.with_suffix(".pdf")
                self._export_pdf(output_docx_path, pdf_path)
                output_pdf_path = str(pdf_path)

            generated.append(
                {
                    "output_docx": str(output_docx_path),
                    "output_pdf": output_pdf_path,
                    "changed_blocks": changed_count,
                    "tables_updated": table_count_updated,
                    "table_rows_generated": rows_generated,
                }
            )

        return {
            "status": "success",
            "action": "generate_documents",
            "rows_total": len(rows_data),
            "generated_count": len(generated),
            "generated": generated,
        }
